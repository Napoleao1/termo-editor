import os
import sys

import json
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt
from PyQt6.QtGui import QIcon, QPixmap
from PyQt6.QtWidgets import (
    QApplication, QWidget, QLabel, QLineEdit, QTextEdit, QPushButton, QVBoxLayout,
    QFileDialog, QScrollArea, QSizePolicy, QMessageBox, QMenu, QComboBox, QDateEdit,
    QGroupBox, QCheckBox, QHBoxLayout
)
from PyQt6.QtCore import Qt, QDate

def resource_path(relative_path):
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

class WordFormApp(QWidget):
    def bloquear_scroll(self, widget):
         widget.wheelEvent = lambda event: event.ignore()

    def __init__(self):
        super().__init__()
        self.setWindowIcon(QIcon(resource_path("logo.ico")))
        self.dark_mode = False
        self.initUI()

    def initUI(self):
        self.setWindowTitle("Preenchimento de Termo")
        self.resize(500, 600)

        scroll = QScrollArea(self)
        scroll.setWidgetResizable(True)
        central_widget = QWidget()
        layout = QVBoxLayout(central_widget)
        self.fields = {}

        self.logo_label = QLabel(self)
        pixmap = QPixmap(resource_path("logo.png"))
        pixmap = pixmap.scaled(170, 170, Qt.AspectRatioMode.KeepAspectRatio, Qt.TransformationMode.SmoothTransformation)
        self.logo_label.setPixmap(pixmap)
        self.logo_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(self.logo_label)

        labels = [
            "Nome", "CPF", "Endereço", "CEP", "Celular", "E-mail",
            "Equipamento", "Patrimônio", "Número de Série", "Observações",
            "Assinatura do Colaborador", "Responsável Técnico"
        ]

        for label in labels:
            layout.addWidget(QLabel(label))
            if label == "Responsável Técnico":
                combo = QComboBox()
                combo.addItems([
                    "", "Roberto Saldanha", "Vitor Kontz", "Ernani Napoleão",
                    "Yuri Loureiro", "Samuel Bispo", "Ricardo Silva"
                ])
                self.bloquear_scroll(combo)
                self.fields[label] = combo
                layout.addWidget(combo)
            elif label == "Observações":
                combo = QComboBox()
                combo.addItems(["", "Retirada", "Devolução", "Desligamento", "Admissão"])
                self.bloquear_scroll(combo)
                self.fields[label] = combo
                layout.addWidget(combo)
            else:
                self.fields[label] = QLineEdit()
                self.fields[label].setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
                self.fields[label].customContextMenuRequested.connect(self.custom_menu)
                layout.addWidget(self.fields[label])

        layout.addWidget(QLabel("Data do Termo"))
        self.fields["Data Completa"] = QDateEdit()
        self.fields["Data Completa"].setCalendarPopup(True)
        self.fields["Data Completa"].setDate(QDate.currentDate())
        self.bloquear_scroll(self.fields["Data Completa"])
        layout.addWidget(self.fields["Data Completa"])

        layout.addWidget(QLabel("Equipamentos Adicionais:"))
        equipamentos = [
            "Mouse sem fio", "Teclado com fio", "Headset Logitech", "Headset JBL",
            "Suporte Notebook", "Mouse pad", "Adaptador Hub USB-C",
            "Carregador/fonte Apple", "Carregador/fonte Dell 65W", "Carregador/fonte Dell 45W"
        ]
        equip_box = QGroupBox()
        equip_layout = QVBoxLayout()
        self.equip_checkboxes = []
        for item in equipamentos:
            cb = QCheckBox(item)
            self.equip_checkboxes.append(cb)
            equip_layout.addWidget(cb)
        equip_box.setLayout(equip_layout)
        layout.addWidget(equip_box)

        layout.addWidget(QLabel("Outro equipamento adicional (manual):"))
        self.fields["Equipamento Extra"] = QLineEdit()
        layout.addWidget(self.fields["Equipamento Extra"])

        btn_row = QHBoxLayout()
        btn_save = QPushButton("💾 Salvar Dados")
        btn_load = QPushButton("📂 Carregar Dados")
        btn_theme = QPushButton("🌙 Tema")
        btn_save.clicked.connect(self.salvar_dados)
        btn_load.clicked.connect(self.carregar_dados)
        btn_theme.clicked.connect(self.trocar_tema)
        btn_row.addWidget(btn_save)
        btn_row.addWidget(btn_load)
        btn_row.addWidget(btn_theme)
        layout.addLayout(btn_row)

        self.btn_generate = QPushButton("Gerar Documento")
        self.btn_generate.clicked.connect(self.generate_doc)
        layout.addWidget(self.btn_generate)

        self.btn_clear = QPushButton("Limpar Campos")
        self.btn_clear.clicked.connect(self.clear_fields)
        layout.addWidget(self.btn_clear)

        footer = QLabel("Powered by Ernani")
        footer.setObjectName("footerLabel")
        footer.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(footer)

        scroll.setWidget(central_widget)
        main_layout = QVBoxLayout(self)
        main_layout.addWidget(scroll)
        self.setLayout(main_layout)
        self.set_light_theme()
    def custom_menu(self, pos):
        sender = self.sender()
        menu = QMenu()
        menu.addAction("Cortar", sender.cut)
        menu.addAction("Copiar", sender.copy)
        menu.addAction("Colar", sender.paste)
        menu.addAction("Selecionar tudo", sender.selectAll)
        menu.exec(sender.mapToGlobal(pos))

    def clear_fields(self):
        for field in self.fields.values():
            if isinstance(field, QLineEdit):
                field.clear()
            elif isinstance(field, QTextEdit):
                field.clear()
            elif isinstance(field, QComboBox):
                field.setCurrentIndex(0)
        for cb in self.equip_checkboxes:
            cb.setChecked(False)

    def salvar_dados(self):
        dados = {}
        for key, field in self.fields.items():
            if isinstance(field, QLineEdit):
                dados[key] = field.text()
            elif isinstance(field, QTextEdit):
                dados[key] = field.toPlainText()
            elif isinstance(field, QComboBox):
                dados[key] = field.currentText()
        dados["Equipamentos Adicionais"] = [cb.text() for cb in self.equip_checkboxes if cb.isChecked()]
        path, _ = QFileDialog.getSaveFileName(self, "Salvar Dados", "", "JSON (*.json)")
        if path:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(dados, f, indent=4, ensure_ascii=False)
            QMessageBox.information(self, "Sucesso", "Dados salvos com sucesso.")

    def carregar_dados(self):
        path, _ = QFileDialog.getOpenFileName(self, "Carregar Dados", "", "JSON (*.json)")
        if path and os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                dados = json.load(f)
                for key, field in self.fields.items():
                    if key in dados:
                        if isinstance(field, QLineEdit):
                            field.setText(dados[key])
                        elif isinstance(field, QTextEdit):
                            field.setPlainText(dados[key])
                        elif isinstance(field, QComboBox):
                            idx = field.findText(dados[key])
                            if idx >= 0:
                                field.setCurrentIndex(idx)
                for cb in self.equip_checkboxes:
                    cb.setChecked(cb.text() in dados.get("Equipamentos Adicionais", []))

    def trocar_tema(self):
        self.dark_mode = not self.dark_mode
        if self.dark_mode:
            self.setStyleSheet("""
                QWidget { background-color: #1e1e1e; color: #ffffff; }
                QLineEdit, QTextEdit, QComboBox, QDateEdit {
                    background-color: #2b2b2b; color: #ffffff;
                    border: 1px solid #6A5ACD;
                }
                QPushButton { background-color: #6A5ACD; color: white; }
                QPushButton:hover { background-color: #483D8B; }
                QCheckBox { color: #ffffff; }
            """)
        else:
            self.set_light_theme()

    def set_light_theme(self):
        self.setStyleSheet("""
            QWidget { background-color: #eaeaea; font-family: 'Segoe UI'; font-size: 12pt; }
            QLabel { font-weight: bold; color: #333; }
            QLineEdit, QTextEdit, QComboBox, QDateEdit {
                border: 1px solid #ccc; border-radius: 5px; padding: 5px;
                background-color: #fff; color: #000;
            }
            QCheckBox {
                color: #000; font-size: 11pt; spacing: 8px;
            }
            QCheckBox::indicator {
                width: 16px; height: 16px;
            }
            QCheckBox::indicator:unchecked {
                border: 2px solid #6A5ACD; background-color: white; border-radius: 3px;
            }
            QCheckBox::indicator:checked {
                border: 2px solid #6A5ACD; background-color: #6A5ACD; border-radius: 3px;
            }
            QComboBox {
                color: #6A5ACD; font-weight: bold;
            }
            QComboBox QAbstractItemView {
                background-color: #fff; color: #6A5ACD; selection-background-color: #dcd0ff;
            }
            QPushButton {
                background-color: #6A5ACD; color: white; font-weight: 600;
                border-radius: 5px; padding: 10px;
            }
            QPushButton:hover {
                background-color: #483D8B;
            }
            QLabel#footerLabel {
                color: #483D8B; font-size: 10pt; font-style: italic; font-weight: bold;
            }
    QScrollBar:vertical {
    background: transparent;
    width: 12px;
    margin: 0px;
}

QScrollBar::handle:vertical {
    background: #6A5ACD; /* roxo */
    min-height: 30px;
    border-radius: 6px;
}

QScrollBar::add-line:vertical,
QScrollBar::sub-line:vertical {
    height: 0;
}

QScrollBar::add-page:vertical,
QScrollBar::sub-page:vertical {
    background: none;
}
QCalendarWidget QWidget {
    color: black;
}

QCalendarWidget QToolButton {
    color: #6A5ACD;  /* cor do mês e ano */
    font-weight: bold;
}

QCalendarWidget QMenu {
    background-color: white;
    color: black;
}


        """)

    def generate_doc(self):
        data = self.fields["Data Completa"].date()
        dia = str(data.day())
        mes_en = data.toString("MMMM").lower()
        meses = {
            "january": "janeiro", "february": "fevereiro", "march": "março", "april": "abril",
            "may": "maio", "june": "junho", "july": "julho", "august": "agosto",
            "september": "setembro", "october": "outubro", "november": "novembro", "december": "dezembro"
        }
        mes_nome = meses.get(mes_en, mes_en)
        mes_num = str(data.month()).zfill(2)
        data_compacta = f"{dia}/{mes_num}/2025"

        file_path = resource_path("Termo responsa.docx")
        if not os.path.exists(file_path):
            QMessageBox.critical(self, "Erro", "Modelo 'Termo responsa.docx' não encontrado.")
            return

        doc = Document(file_path)

        for p in doc.paragraphs:
            if "[INSERIR NOME]" in p.text:
                p.clear()
                run = p.add_run("Eu, ")
                run.font.bold = False
                run = p.add_run(self.fields["Nome"].text())
                run.font.bold = True
                run = p.add_run(", inscrito no CPF sob nº ")
                run.font.bold = False
                run = p.add_run(self.fields["CPF"].text())
                run.font.bold = True
                run = p.add_run(", residente e domiciliado na ")
                run.font.bold = False
                endereco = f"{self.fields['Endereço'].text()}"
                cep = self.fields['CEP'].text().strip()
                if cep:
                    endereco += f" - CEP: {cep}"
                run = p.add_run(endereco)
                run.font.bold = True

                celular = self.fields["Celular"].text().strip()
                if celular:
                    run = p.add_run(", celular: ")
                    run.font.bold = False
                    run = p.add_run(celular)
                    run.font.bold = True

                email = self.fields["E-mail"].text().strip()
                if email:
                    run = p.add_run(", Email pessoal: ")
                    run.font.bold = False
                    run = p.add_run(email)
                    run.font.bold = True
            else:
                p.text = p.text.replace("[INSERIR NOME]", self.fields["Nome"].text())
                p.text = p.text.replace("[INSERIR CPF]", self.fields["CPF"].text())
                endereco = f"{self.fields['Endereço'].text()} - CEP: {self.fields['CEP'].text()}"
                p.text = p.text.replace("[INSERIR ENDEREÇO COMPLETO COM CEP]", endereco)
                p.text = p.text.replace("[INSERIR NÚMERO DO CELULAR]", self.fields["Celular"].text())
                p.text = p.text.replace("[INSERIR E-MAIL PESSOAL]", self.fields["E-mail"].text())
                p.text = p.text.replace("[INSERIR DATA]", data_compacta)
                p.text = p.text.replace("[DIA]", dia)
                p.text = p.text.replace("[MÊS]", mes_nome)
                p.text = p.text.replace("[ANO]", "2025")

                if "Assinatura do colaborador:" in p.text:
                    assinatura = self.fields["Assinatura do Colaborador"].text().strip()
                    if assinatura:
                        p.text = p.text.replace("Assinatura do colaborador:", f"Assinatura do colaborador: {assinatura}")
                    else:
                        p.text = ""

                if "Responsável técnico:" in p.text:
                    tecnico = self.fields["Responsável Técnico"].currentText().strip()
                    if tecnico:
                        p.text = p.text.replace("Responsável técnico:", f"Responsável técnico: {tecnico}")
                    else:
                        p.text = ""

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        text = text.replace("Equipamento:", f"Equipamento: {self.fields['Equipamento'].text()}")
                        text = text.replace("Patrimônio:", f"Patrimônio: {self.fields['Patrimônio'].text()}")
                        text = text.replace("Número de série:", f"Número de série: {self.fields['Número de Série'].text()}")

                        adicionais = [cb.text() for cb in self.equip_checkboxes if cb.isChecked()]
                        extra = self.fields["Equipamento Extra"].text().strip()
                        if extra:
                            adicionais.append(extra)
                       
                        if adicionais:
                            text = text.replace("Equipamentos Adicionais:", f"Equipamentos Adicionais: " + ", ".join(adicionais))


                        obs = self.fields["Observações"].currentText().strip()
                        text = text.replace("Observações:", f"Observações: {obs}")
                        paragraph.clear()
                        paragraph.add_run(text).font.size = Pt(9)

        save_path, _ = QFileDialog.getSaveFileName(self, "Salvar Documento", "", "Arquivos Word (*.docx)")
        if save_path:
            doc.save(save_path)
            reply = QMessageBox.question(self, 'Salvar como PDF?', 'Deseja salvar também como PDF?',
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            if reply == QMessageBox.StandardButton.Yes:
                self.save_as_pdf(save_path)

    def save_as_pdf(self, word_path):
        pdf_output = word_path.replace(".docx", ".pdf")
        libreoffice_path = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
        if not os.path.exists(libreoffice_path):
            QMessageBox.critical(self, "Erro", "LibreOffice não encontrado.")
            return
        try:
            subprocess.run([
                libreoffice_path, "--headless", "--convert-to", "pdf", word_path,
                "--outdir", os.path.dirname(pdf_output)
            ], check=True, capture_output=True, text=True)
            QMessageBox.information(self, "PDF Salvo", f"Arquivo PDF salvo em: {pdf_output}")
        except subprocess.CalledProcessError as e:
            QMessageBox.critical(self, "Erro", f"Erro ao converter para PDF:\n{e.stderr}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = WordFormApp()
    window.show()
    sys.exit(app.exec())


